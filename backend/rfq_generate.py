"""
Генерация финального RFQ (.docx) из полей, сопоставленных на фронтенде.

Работает поверх реального Word-шаблона (templates/rfq_template.docx) —
заполняются только значения ячеек, форматирование (шрифты, цвета,
границы таблиц) остаётся тем, что задано в шаблоне. Никакого ИИ здесь
нет — просто подстановка текста в конкретные ячейки по их положению
в фиксированной структуре шаблона.
"""

import os
from copy import copy

import docx
from docx.shared import Pt

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "rfq_template.docx")

# Индекс строки в первой таблице шаблона для каждого поля RFQ.
FIELD_ROW = {
    "issuedBy": 0,
    "invitedBidder": 1,
    "rfqReference": 2,
    "issueDate": 3,
    "workLocation": 4,
    "expectedStart": 5,
    "requiredCompletion": 6,
    "bidDate": 7,
    "projectName": 10,
    "projectDescription": 11,
    "scope": 12,
}


def _set_value(cell, text):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if not text:
        return
    lines = str(text).split("\n")
    run = p.add_run(lines[0])
    run.font.name = "Century Gothic"
    run.font.size = Pt(10)
    for extra in lines[1:]:
        new_p = cell.add_paragraph()
        r = new_p.add_run(extra)
        r.font.name = "Century Gothic"
        r.font.size = Pt(10)


def generate_rfq_docx(fields: dict, output_path: str) -> str:
    d = docx.Document(TEMPLATE_PATH)
    table = d.tables[0]

    for field_id, row_idx in FIELD_ROW.items():
        value = fields.get(field_id, "")
        _set_value(table.rows[row_idx].cells[1], value)

    d.save(output_path)
    return output_path
